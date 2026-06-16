"""Generare dosar per entitate (DOCX în memorie) pentru SOLOMONAR.

DOCX pur (python-docx) — fără dependențe de sistem; se deschide în Word/Google Docs și se
exportă în PDF dintr-un click. Folosit de fișa Persoane (buton de download).
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

PURPLE = RGBColor(0x6D, 0x28, 0xD9)
DARK = RGBColor(0x1C, 0x1E, 0x26)
GREY = RGBColor(0x55, 0x5B, 0x66)


def _lei(n) -> str:
    try:
        n = float(n)
    except (TypeError, ValueError):
        return "—"
    for div, suf in ((1e9, " mld"), (1e6, " mil"), (1e3, " mii")):
        if abs(n) >= div:
            return f"{n/div:.1f}{suf} lei".replace(".", ",")
    return f"{n:,.0f} lei".replace(",", ".")


def _int(n) -> str:
    try:
        return f"{int(n):,}".replace(",", ".")
    except (TypeError, ValueError):
        return "—"


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = PURPLE


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cp = cells[i].paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            r = cp.add_run(str(val))
            r.font.size = Pt(9)
    if widths:
        for col, w in enumerate(widths):
            for cell in t.columns[col].cells:
                cell.width = Cm(w)
    return t


def persoana_docx(p: dict, generat: str = "") -> bytes:
    """Construiește un dosar DOCX pentru o persoană (dict din data.persoana) -> bytes."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6); sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.9); sec.right_margin = Cm(1.9)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(10); st.font.color.rgb = DARK

    nume = (p.get("nume_key") or "").title()
    pl = p.get("parlamentar") or {}

    # ---- antet ----
    h = doc.add_paragraph(); h.paragraph_format.space_after = Pt(0)
    r = h.add_run(nume); r.bold = True; r.font.size = Pt(20)
    sub = doc.add_paragraph(); sub.paragraph_format.space_after = Pt(8)
    bits = [f"nivel de încredere a legăturilor: {p.get('incredere', '—')}"]
    for k in ("camera", "partid", "judet"):
        if pl.get(k):
            bits.append(str(pl[k]))
    rr = sub.add_run(" · ".join(bits)); rr.font.size = Pt(10); rr.font.color.rgb = GREY

    # ---- sumar ----
    s = doc.add_paragraph()
    summ = (f"Declarații: {_int(p.get('n_declaratii', 0))}   ·   "
            f"Companii conduse: {_int(p.get('n_companii', 0))}   ·   "
            f"Contracte de stat: {_lei(p.get('total_contracte_ron') or 0)}   ·   "
            f"Achiziții directe: {_lei(p.get('total_achizitii_directe_ron') or 0)}")
    r = s.add_run(summ); r.bold = True; r.font.size = Pt(10)

    # ---- mandat ----
    if pl:
        _heading(doc, "Mandat parlamentar")
        md = [("Cameră", pl.get("camera") or "—"), ("Partid", pl.get("partid") or "—"),
              ("Județ", pl.get("judet") or "—"),
              ("Comisii", _int(len(pl.get("comisii") or []))),
              ("Proiecte inițiate (PLx)", _int(pl.get("plx_initiate") or 0))]
        _table(doc, ["Câmp", "Valoare"], md, widths=[5, 11])

    # ---- declarații ----
    decl = p.get("declaratii") or []
    if decl:
        _heading(doc, f"Declarații de avere / interese ({len(decl)})")
        rows = [((d.get("tip") or "").title(), d.get("institutie") or "—",
                 _lei(d.get("venituri_ron")) if d.get("venituri_ron") else "—") for d in decl[:40]]
        _table(doc, ["Tip", "Instituție", "Venituri"], rows, widths=[2.5, 10, 3.5])

    # ---- companii conduse ----
    comp = p.get("companii") or []
    if comp:
        _heading(doc, f"Companii conduse ({len(comp)})")
        rows = []
        for c in sorted(comp, key=lambda x: -((x.get("contracte_stat") or {}).get("total_ron") or 0)):
            cs = c.get("contracte_stat") or {}
            rows.append((c.get("nume") or "—", c.get("rol") or "—",
                         c.get("sector") or "—", _lei(cs.get("total_ron")) if cs.get("total_ron") else "—"))
        _table(doc, ["Firmă", "Rol", "Sector", "Contracte stat"], rows[:40], widths=[7, 3.5, 3, 3])

    # ---- conflicte autodeclarate ----
    fa = p.get("firme_contracte_autodeclarate") or []
    if fa:
        _heading(doc, "⚠ Firme autodeclarate cu contracte de stat (conflict documentat)")
        rows = [(f.get("nume") or "—", str(f.get("cui") or "—"),
                 _lei(f.get("total_ron")) if f.get("total_ron") else "—") for f in fa[:20]]
        _table(doc, ["Firmă", "CUI", "Contracte"], rows, widths=[8, 3.5, 4.5])

    # ---- CV ----
    cv = p.get("cv") or {}
    if isinstance(cv, dict) and (cv.get("studii") or cv.get("experienta")):
        _heading(doc, "CV — studii & experiență")
        for lbl, key in (("Studii", "studii"), ("Experiență", "experienta")):
            txt = (cv.get(key) or "").strip()
            if txt:
                pp = doc.add_paragraph(); pp.paragraph_format.space_after = Pt(3)
                rr = pp.add_run(lbl + ": "); rr.bold = True; rr.font.size = Pt(9)
                rr = pp.add_run(txt[:1500]); rr.font.size = Pt(9)

    # ---- footer / disclaimer ----
    doc.add_paragraph()
    f = doc.add_paragraph()
    r = f.add_run(
        "Dosar generat automat de SOLOMONAR din date publice agregate. Legăturile persoană↔firmă "
        "fără CNP se fac pe nume — pot exista omonimi (vezi nivelul de încredere). Comunicatele DNA "
        "sunt trimiteri în judecată, nu condamnări. Acesta este un punct de plecare pentru verificare, "
        "nu un verdict." + (f"  ·  Generat: {generat}" if generat else ""))
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
